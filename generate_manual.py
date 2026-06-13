"""Generate PMS User Manual as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.2)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)   # blue-700

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold  = bold
            run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run("Tip: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x05, 0x6F, 0x00)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # light green shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E6F4EA')
    pPr.append(shd)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x92, 0x4E, 0x00)
    r2 = p.add_run(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8E1')
    pPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # blue fill
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E40AF')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Photo Management System")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PMS")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x60, 0x7D, 0x8B)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("User Manual")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Version 1.0  |  {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual – auto TOC requires Word field updates)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_entries = [
    ("1", "Getting Started", "4"),
    ("1.1", "System Requirements", "4"),
    ("1.2", "Download & Install", "4"),
    ("1.3", "First Launch", "5"),
    ("2", "Understanding the Interface", "6"),
    ("2.1", "Main Layout Overview", "6"),
    ("2.2", "Sidebar Tabs", "6"),
    ("2.3", "Main Content Panel", "7"),
    ("3", "Managing Folders (Tree)", "7"),
    ("3.1", "Creating Folders", "7"),
    ("3.2", "Renaming and Deleting Folders", "8"),
    ("3.3", "Reordering and Moving Folders", "8"),
    ("3.4", "Serial Numbers (SN)", "9"),
    ("4", "Working with Documents", "9"),
    ("4.1", "Uploading Files", "9"),
    ("4.2", "Uploading a Folder", "10"),
    ("4.3", "Linking Existing Documents", "10"),
    ("4.4", "Viewing a Document", "10"),
    ("4.5", "Downloading and Deleting Documents", "11"),
    ("5", "OCR – Extract Text from Documents", "11"),
    ("5.1", "How OCR Works", "11"),
    ("5.2", "Running OCR", "12"),
    ("5.3", "Saving OCR Text for Search", "12"),
    ("6", "Key Parameter Extraction", "13"),
    ("6.1", "What Are Key Parameters?", "13"),
    ("6.2", "Extracting Parameters from a Document", "13"),
    ("6.3", "Managing Global Key Parameters", "14"),
    ("7", "Searching Documents", "14"),
    ("7.1", "Basic Search", "14"),
    ("7.2", "Using Filters", "15"),
    ("7.3", "Opening and Jumping to Results", "15"),
    ("8", "Generating a Databook PDF", "15"),
    ("8.1", "Selecting Documents", "15"),
    ("8.2", "Generating the PDF", "16"),
    ("9", "Merging Documents", "16"),
    ("10", "Project Management", "17"),
    ("10.1", "Export Project (.dms)", "17"),
    ("10.2", "Import Project (.dms)", "17"),
    ("10.3", "CSV Export", "18"),
    ("10.4", "Batch ZIP Import", "18"),
    ("10.5", "Hierarchy Import", "18"),
    ("11", "Password Protection", "19"),
    ("12", "Keyboard Shortcuts", "20"),
    ("13", "Troubleshooting", "20"),
]
for num, title, page in toc_entries:
    p = doc.add_paragraph()
    indent = Inches(0.4) if "." in num else Inches(0)
    p.paragraph_format.left_indent = indent
    run = p.add_run(f"{num}  {title}")
    if "." not in num:
        run.bold = True
    tab_stop = p.paragraph_format
    p.add_run(f"\t{page}")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1  GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
h1("1  Getting Started")

h2("1.1  System Requirements")
add_table(
    ["Item", "Requirement"],
    [
        ["Operating System", "Windows 10 / 11 (64-bit)"],
        ["RAM", "4 GB minimum, 8 GB recommended"],
        ["Disk Space", "200 MB for the app; additional space for your documents"],
        ["Display", "1280 × 720 or larger"],
        ["Internet", "Not required (runs fully offline). Internet needed only for GPS address lookup."],
    ],
    col_widths=[1.8, 4.2]
)

h2("1.2  Download & Install")
para("Follow these steps to get PMS running on your Windows computer.")
numbered("Receive or download the file PMS.zip from your contact (e.g. email, WeChat, USB drive).")
numbered("Right-click PMS.zip and choose Extract All…")
numbered("Choose a destination folder, for example C:\\PMS, and click Extract.")
numbered("Open the extracted folder. You will see a file named PMS.exe.")
numbered("Double-click PMS.exe to launch the application.")
tip("You do not need to install anything. PMS.exe is a standalone program – it includes Python and all libraries.")
note("Windows may show a SmartScreen warning the first time you run PMS.exe. Click More info → Run anyway. This warning appears because the file is not code-signed; it is safe to proceed.")
doc.add_paragraph()

h2("1.3  First Launch")
para("When PMS starts for the first time:")
numbered("A small console window opens briefly in the background – this is the built-in web server. Do not close it.")
numbered("Your default web browser opens automatically to http://localhost:5000.")
numbered("You will see the PMS welcome screen.")
numbered("You will be prompted to choose a Storage Folder – this is where all your files and data will be saved.")
bullet("Click the folder icon in the header bar (or the blue Set Storage Path button).")
bullet("Type a path such as C:\\Users\\YourName\\Documents\\PMS_Data and press Enter.")
bullet("PMS creates the folder automatically if it does not exist.")
numbered("Once a storage path is set, the main interface loads and you are ready to use PMS.")
tip("Choose a storage path inside your Documents or Pictures folder so it is included in your regular backups.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2  INTERFACE
# ══════════════════════════════════════════════════════════════════════════════
h1("2  Understanding the Interface")

h2("2.1  Main Layout Overview")
para("The PMS window is divided into three main areas:")
add_table(
    ["Area", "Location", "Purpose"],
    [
        ["Header Bar", "Top of the page", "Shows storage path, total document count, Project menu, and logout button."],
        ["Sidebar", "Left panel", "Tabs for Tree view, Search, and Databook generation."],
        ["Main Content Panel", "Right / centre area", "Shows documents belonging to the currently selected folder."],
    ],
    col_widths=[1.5, 1.8, 3.3]
)
tip("Drag the vertical divider between the sidebar and the content panel to resize. Double-click the divider to reset to default width.")

h2("2.2  Sidebar Tabs")
add_table(
    ["Tab", "Chinese Label", "What It Does"],
    [
        ["Tree", "树形", "Browse and manage your folder hierarchy. This is the primary navigation view."],
        ["Search", "搜索", "Full-text and metadata search across all documents."],
        ["Databook", "数据手册", "Select documents and generate a professional combined PDF with cover and table of contents."],
    ],
    col_widths=[1.0, 1.4, 4.2]
)

h2("2.3  Main Content Panel")
para("When you click a folder in the Tree tab, the right panel shows:")
bullet("Folder name, serial number (SN), and description at the top.")
bullet("Action buttons: Upload Files, Upload Folder, Link Existing Document, Merge Documents.")
bullet("Cards for each document in the folder – showing file name, size, type icon, upload date, and action icons.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3  FOLDERS
# ══════════════════════════════════════════════════════════════════════════════
h1("3  Managing Folders (Tree)")
para("PMS organises your documents into a tree of folders – similar to Windows Explorer but with extra features like serial numbers, descriptions, and document counts.")

h2("3.1  Creating Folders")
numbered("Click a folder in the tree to select it (the new folder will be created as a child of the selected folder).")
numbered("Click the New Folder button (＋ icon) in the sidebar toolbar, OR right-click the selected folder and choose New Folder.")
numbered("A text field appears. Type the folder name and press Enter to confirm, or press Escape to cancel.")
tip("You can nest folders as many levels deep as you need. There is no limit.")

h2("3.2  Renaming and Deleting Folders")
para("Rename:")
bullet("Double-click the folder name, or right-click → Rename.")
bullet("Edit the name and press Enter.")

para("Delete:")
bullet("Right-click the folder → Delete.")
bullet("Confirm the deletion. The folder and all its sub-folders are removed from the tree.")
note("Deleting a folder removes it from the tree, but the physical document files remain on your disk inside the storage folder. You can still find them via the Search tab or by browsing the storage folder in Windows Explorer.")

h2("3.3  Reordering and Moving Folders")
para("Reorder siblings (folders at the same level):")
bullet("Hover over a folder to reveal ▲ and ▼ arrow buttons.")
bullet("Click ▲ to move the folder up, ▼ to move it down among its siblings.")

para("Move to a different parent:")
bullet("Click and drag the folder.")
bullet("Drop it onto the target parent folder (a highlight shows the drop target).")
bullet("You can also drop it just above or below another folder to place it before or after that folder.")

h2("3.4  Serial Numbers (SN)")
para("Every folder can have an optional Serial Number (SN) used to tag assemblies, parts, or project identifiers.")
bullet("Click the pencil icon on a folder to open its properties.")
bullet("Enter a value in the Serial Number field (e.g. SN-YT635R0267).")
bullet("The SN is displayed next to the folder name in the tree.")
bullet("Documents inside the folder automatically inherit the folder's SN unless they have their own SN override.")
tip("To give a specific document its own SN, click the SN button on its document card.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4  DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4  Working with Documents")

h2("4.1  Uploading Files")
numbered("Select the destination folder in the Tree tab.")
numbered("Click the Upload Files button in the main content panel.")
numbered("In the file picker, select one or more files and click Open.")
numbered("A progress bar appears at the top of the screen showing X / Y files uploaded.")
numbered("When complete, the new document cards appear in the panel.")

para("Supported file types:")
add_table(
    ["Category", "Extensions"],
    [
        ["Images", "JPEG, PNG, GIF, WebP, BMP, TIFF, HEIC"],
        ["Documents", "PDF, DOC, DOCX, XLS, XLSX, PPT, PPTX, TXT, CSV, HTML, JSON, XML"],
        ["Archives", "ZIP"],
    ],
    col_widths=[1.5, 5.0]
)
tip("For JPEG and TIFF photos, PMS automatically reads the camera date and GPS location from the EXIF metadata and stores it with the document.")

h2("4.2  Uploading a Folder")
numbered("Select the destination folder in the tree.")
numbered("Click the Upload Folder button.")
numbered("Select a folder on your computer.")
numbered("PMS uploads all files recursively and creates matching sub-folders in the tree.")

h2("4.3  Linking Existing Documents")
para("The same document can appear in multiple folders without duplicating the file on disk.")
numbered("Select the folder you want to link the document to.")
numbered("Click Link Existing Document.")
numbered("A dialog lists all documents in the system.")
numbered("Check the box next to the document(s) you want and click Add Link.")
numbered("The document now appears in both the original folder and the new folder.")

h2("4.4  Viewing a Document")
numbered("Click the document name or thumbnail to open the Document Viewer.")
numbered("The viewer opens as an overlay showing a preview on the left.")
bullet("Images are shown inline.")
bullet("PDFs open in an embedded viewer with scroll and zoom.")
bullet("Text files are shown in a readable monospace view.")
bullet("Office files (Word, Excel, PowerPoint) offer a Download button to open in your local Office app.")
numbered("The right side of the viewer has buttons for OCR, Key Parameter Extraction, and Download.")
numbered("Click the ✕ button or press Escape to close the viewer.")

h2("4.5  Downloading and Deleting Documents")
para("Download: Click the download icon (⬇) on any document card, or click Download inside the viewer.")

para("Unlink from folder: Click the trash icon (🗑) on the document card. If the document is linked to other folders, it is only removed from this folder – the file and other links remain intact.")

para("Delete permanently: Only available from the Search tab or the Unlinked Documents list. Permanently removes the document from all folders and deletes the file from disk.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5  OCR
# ══════════════════════════════════════════════════════════════════════════════
h1("5  OCR – Extract Text from Documents")

h2("5.1  How OCR Works")
para("OCR (Optical Character Recognition) converts scanned images and PDFs into searchable text. PMS uses the Tesseract OCR engine, which supports both English and Chinese (Simplified).")
para("PMS uses two methods to extract text:")
add_table(
    ["Method", "When Used", "Speed"],
    [
        ["Text Layer", "PDFs that already contain digital text", "Very fast"],
        ["Tesseract OCR", "Scanned PDFs, images, or corrupted text layers", "Moderate (depends on file size)"],
    ],
    col_widths=[1.5, 3.5, 1.5]
)
note("If Tesseract is not installed on the server machine, OCR for images will not be available. The text layer extraction for digital PDFs will still work.")

h2("5.2  Running OCR")
numbered("Open the document viewer by clicking the document name.")
numbered("Click Extract Text (OCR) in the right panel.")
numbered("PMS extracts text and displays it in the text panel on the right.")
numbered("If the document is a PDF with an embedded text layer, extraction is instant.")
numbered("If OCR is needed (scanned document or image), a progress indicator appears.")
numbered("When finished, the extracted text is shown in an editable text area.")
bullet("You may correct OCR errors by editing the text directly.")
bullet("The panel also shows: number of pages processed, character count, language detected, and extraction method used.")
note("If a PDF has a corrupted or incorrect text layer, click Force OCR to re-extract using Tesseract instead.")

h2("5.3  Saving OCR Text for Search")
numbered("After reviewing (and optionally editing) the extracted text, click Save & Make Searchable.")
numbered("The text is saved with the document's metadata.")
numbered("The document now appears in Search results when you search for words it contains.")
tip("You only need to run OCR once per document. The text is stored and the document remains searchable forever.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6  KEY PARAMETERS
# ══════════════════════════════════════════════════════════════════════════════
h1("6  Key Parameter Extraction")

h2("6.1  What Are Key Parameters?")
para("Key Parameters are user-defined labels (like 'Contract Number', 'Pressure Rating', '人物', '事件') that PMS can automatically find and extract from document text. Each parameter stores two values:")
bullet("Design value – the required or specified value (e.g. '3000 psi').")
bullet("Actual value – the measured or certified value found in the document.")
para("This lets you quickly compare specifications against actual values across hundreds of documents.")

h2("6.2  Extracting Parameters from a Document")
numbered("Open the document viewer.")
numbered("Click Extract Key Parameters in the right panel.")
numbered("A list of parameters appears (global defaults + any document-specific ones).")
numbered("Check the parameters you want to extract.")
numbered("Optionally toggle Same Line Only to only match values on the same line as the parameter name.")
numbered("Click Extract.")
numbered("PMS searches the document text and fills in a results table:")
add_table(
    ["Column", "Meaning"],
    [
        ["Parameter", "The label being searched for"],
        ["Found Value", "The value PMS found in the document text"],
        ["Status", "✓ Found or NF (Not Found)"],
    ],
    col_widths=[1.5, 5.0]
)
numbered("Edit any values if needed.")
numbered("Click Save Parameters to store the values in the document's metadata.")
tip("Metadata is global – changes apply everywhere this document is referenced.")

h2("6.3  Managing Global Key Parameters")
numbered("Click the Project menu in the header.")
numbered("Choose Global Key Parameters.")
numbered("In the dialog, add, remove, or reorder parameters.")
numbered("Click Restore Defaults to reset to the default Chinese list.")
numbered("Changes take effect immediately for all future extractions.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7  SEARCH
# ══════════════════════════════════════════════════════════════════════════════
h1("7  Searching Documents")

h2("7.1  Basic Search")
numbered("Click the Search tab (搜索) in the sidebar.")
numbered("Type a word or phrase in the search box.")
numbered("Results appear instantly as cards showing:")
bullet("Document name and file type icon.")
bullet("Upload date, file size, document ID.")
bullet("All metadata fields (parameter name / design value / actual value).")
bullet("GPS-derived location (if available).")
bullet("Which folders reference this document.")
note("Search is case-insensitive and partial-match – typing 'pump' will find 'Pump Assembly', 'centrifugal pump', etc.")

h2("7.2  Using Filters")
add_table(
    ["Filter", "How to Use"],
    [
        ["File Type", "Check Image, PDF, Text, or Documents to restrict results to that type."],
        ["Scope Node", "Enable Scope to selected node to search only within the currently selected folder and its sub-folders."],
        ["Metadata", "Enter a parameter name and/or value to find documents with matching metadata (e.g. Pressure = 3000 psi)."],
    ],
    col_widths=[1.5, 5.0]
)

h2("7.3  Opening and Jumping to Results")
bullet("Click a result card to open the Document Viewer for that document.")
bullet("Double-click a result card to switch to the Tree tab and jump directly to the folder containing that document.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8  DATABOOK
# ══════════════════════════════════════════════════════════════════════════════
h1("8  Generating a Databook PDF")
para("The Databook feature combines selected documents into a single, professional PDF with a cover page, table of contents, bookmarks, section headers, and page numbers.")

h2("8.1  Selecting Documents")
numbered("Click the Databook tab (数据手册) in the sidebar.")
numbered("The tree is shown with a checkbox next to each document.")
numbered("Check individual documents, or:")
bullet("Click Select All at Node to select all documents in a specific folder.")
bullet("Click Select All in Subtree to select all documents in a folder and all its sub-folders.")
bullet("Click Clear to deselect all.")
tip("Your selection is saved automatically, so you can close and reopen PMS and your selection will still be there.")

h2("8.2  Generating the PDF")
numbered("After selecting documents, click Generate Databook.")
numbered("In the dialog, enter:")
bullet("Title – the main title shown on the cover page.")
bullet("Subtitle – a secondary line (e.g. project name, revision number).")
numbered("Click Generate.")
numbered("PMS creates the PDF on the server and your browser downloads it automatically.")

para("The generated PDF includes:")
bullet("Cover page with title, subtitle, and date.")
bullet("Table of Contents with clickable bookmarks.")
bullet("Section header pages matching your folder names.")
bullet("All selected documents in tree order.")
bullet("Page numbers and footer on every page.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9  MERGE DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("9  Merging Documents")
para("Merge combines multiple PDFs (and images) from a folder into a single PDF file, which is then saved back to that folder.")
numbered("Select a folder in the Tree tab.")
numbered("Click the Merge Documents button in the main content panel.")
numbered("A dialog opens showing all documents in the folder.")
numbered("Optionally check Include Subfolders to also include documents from sub-folders.")
numbered("Check the documents you want to merge (or click Select All).")
numbered("Enter a name for the output file in the Output Filename field.")
numbered("Click Merge.")
numbered("PMS creates the merged PDF and adds it to the current folder.")
note("The original documents remain in the folder after merging. You can manually unlink them if you no longer need individual copies.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10  PROJECT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("10  Project Management")
para("The Project menu (in the header bar) provides tools for backing up, restoring, and importing data.")

h2("10.1  Export Project (.dms)")
para("Creates a complete backup of your entire project – all folder structure, document metadata, and document files – as a single .dms file (which is a ZIP archive).")
numbered("Click Project in the header bar.")
numbered("Choose Export Project (.dms).")
numbered("Your browser downloads a file named something like project_backup_20260610.dms.")
tip("Save this file to a USB drive or cloud storage as a regular backup.")

h2("10.2  Import Project (.dms)")
para("Restores a previously exported project on any computer running PMS.")
numbered("Click Project → Import Project (.dms).")
numbered("Select the .dms file.")
numbered("In the dialog, choose a Target Folder on your computer (a new empty folder is recommended).")
numbered("Click Import.")
numbered("PMS extracts all files and rebuilds the tree structure.")
note("If you import into a folder that already contains data, you will be prompted to confirm whether to overwrite.")

h2("10.3  CSV Export")
para("Exports all document metadata as a CSV spreadsheet.")
numbered("Click Project → Export CSV.")
numbered("Your browser downloads a CSV file.")
para("The CSV contains one row per document with columns for: Document ID, Name, File Size, Upload Date, and all metadata parameter fields.")
tip("Open the CSV in Microsoft Excel or Google Sheets for analysis, reporting, or sharing with colleagues.")

h2("10.4  Batch ZIP Import")
para("Quickly upload many files and automatically sort them into folders using a naming convention.")
numbered("Prepare a ZIP file. Name each file using the pattern: FolderName#Description.pdf (e.g. Pump_Assembly#test_report.pdf).")
numbered("Click Project → Batch ZIP Import.")
numbered("Select your ZIP file.")
numbered("PMS reads each filename, matches it to the folder name before the # character, and places the document in that folder.")
note("Folders must already exist in the tree before importing. Documents with unrecognised folder names are placed in an Unlinked state.")

h2("10.5  Hierarchy Import (Build Tree from Text File)")
para("Quickly create a folder tree structure from a plain text file.")
numbered("Prepare a text file (.txt). Each line should contain: NodeID [ParentNodeID]")
para("Example:")
p = doc.add_paragraph("Assembly\nPump_Housing Assembly\nRotor Assembly\nImpeller Rotor\nShaft Rotor")
p.style = doc.styles['No Spacing']
p.paragraph_format.left_indent = Inches(0.4)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
doc.add_paragraph()
numbered("Click Project → Create Folder Tree from Hierarchy.")
numbered("Upload your text file (or paste the text).")
numbered("PMS validates the structure, shows a preview, and reports how many nodes will be created.")
numbered("Click Create.")
numbered("All folders appear in the tree immediately.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11  PASSWORD
# ══════════════════════════════════════════════════════════════════════════════
h1("11  Password Protection")
para("You can protect PMS with a password so that only authorised users can access the data.")

para("Set a password:")
numbered("Click Project → Set Password.")
numbered("Enter and confirm your chosen password.")
numbered("Click Set.")
numbered("On the next PMS startup, a login screen will appear before the main interface.")

para("Remove the password:")
numbered("Click Project → Change Password.")
numbered("Enter your current password to confirm, then leave the new password fields blank.")
numbered("Click Set to remove password protection.")

para("Logout:")
bullet("Click the Logout button in the header bar.")
bullet("PMS returns to the login screen. The server keeps running; only the browser session ends.")

note("The password is stored as a salted hash on the server using 100,000 iterations of SHA-256 – it is not stored in plain text.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12  KEYBOARD SHORTCUTS
# ══════════════════════════════════════════════════════════════════════════════
h1("12  Keyboard Shortcuts")
add_table(
    ["Action", "Shortcut"],
    [
        ["Confirm folder rename / new folder name", "Enter"],
        ["Cancel rename / new folder", "Escape"],
        ["Start renaming a folder", "Double-click the folder name"],
        ["Open folder context menu", "Right-click the folder"],
        ["Close the Document Viewer", "Escape  or click  ✕"],
        ["Copy selected text", "Ctrl + C"],
        ["Paste text", "Ctrl + V"],
        ["Select all text in a text field", "Ctrl + A"],
        ["Undo (in text fields)", "Ctrl + Z"],
    ],
    col_widths=[3.5, 3.0]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13  TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
h1("13  Troubleshooting")

add_table(
    ["Problem", "Solution"],
    [
        [
            "Windows SmartScreen blocks PMS.exe",
            "Click 'More info' then 'Run anyway'. This appears because the file is not commercially signed; it is safe."
        ],
        [
            "The browser does not open automatically",
            "Open your browser manually and go to http://localhost:5000"
        ],
        [
            "Browser shows 'This site can't be reached'",
            "Make sure the PMS.exe console window is still running in the background. Relaunch PMS.exe if needed."
        ],
        [
            "OCR button is greyed out or says Tesseract not available",
            "Tesseract OCR is an optional component. Install it from https://github.com/UB-Mannheim/tesseract/wiki and restart PMS."
        ],
        [
            "Chinese text is not extracted correctly",
            "Install the Tesseract Chinese Simplified language pack (chi_sim.traineddata). See Tesseract documentation."
        ],
        [
            "GPS location shows as coordinates, not a city name",
            "The reverse-geocoding service (Nominatim) requires internet access. Check your network connection."
        ],
        [
            "Uploaded file does not appear in the document list",
            "Refresh the browser page (F5). The file may have been uploaded successfully but the page did not update."
        ],
        [
            "Cannot find a document I uploaded earlier",
            "Use the Search tab to search for the filename or content. The document may be in a different folder."
        ],
        [
            "PMS.exe closes immediately on launch",
            "Make sure no other program is using port 5000 on your computer. Restart your computer and try again."
        ],
        [
            "Databook PDF generation fails",
            "Ensure all selected documents are valid PDFs or images. Corrupted files can prevent generation."
        ],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_paragraph()
h2("Getting Help")
para("If you encounter an issue not listed above, please contact the application developer with:")
bullet("A description of what you were doing when the problem occurred.")
bullet("Any error message shown on screen.")
bullet("Your operating system version (e.g. Windows 11).")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
h1("Quick Reference Card")
add_table(
    ["Task", "How"],
    [
        ["Create a new folder", "Select parent → click ＋ New Folder → type name → Enter"],
        ["Upload documents", "Select folder → Upload Files → pick files"],
        ["View a document", "Click the document name"],
        ["Run OCR on a document", "Open viewer → Extract Text (OCR) → Save & Make Searchable"],
        ["Search all documents", "Click Search tab → type keyword"],
        ["Export project backup", "Project menu → Export Project (.dms)"],
        ["Generate a combined PDF", "Databook tab → check documents → Generate Databook"],
        ["Merge files into one PDF", "Select folder → Merge Documents → select files → Merge"],
        ["Move a folder", "Drag & drop the folder to new location in tree"],
        ["Set a serial number", "Click pencil icon on folder → enter SN → Done"],
        ["Export metadata as CSV", "Project menu → Export CSV"],
        ["Set a password", "Project menu → Set Password"],
    ],
    col_widths=[2.5, 4.5]
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/david/PMS/PMS_User_Manual.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
