# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.2)
section.top_margin  = section.bottom_margin = Inches(1)

FONT = 'Microsoft YaHei'
BLUE1 = RGBColor(0x1E, 0x40, 0xAF)
BLUE2 = RGBColor(0x1D, 0x4E, 0x89)

def _cn(run):
    run.font.name = FONT
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    except Exception:
        pass

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: _cn(r); r.font.color.rgb = BLUE1

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: _cn(r); r.font.color.rgb = BLUE2

def para(text):
    p = doc.add_paragraph()
    r = p.add_run(text); _cn(r)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text); _cn(r)

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text); _cn(r)

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = p.paragraph_format.right_indent = Inches(0.4)
    r1 = p.add_run('提示：'); _cn(r1); r1.bold = True; r1.font.color.rgb = RGBColor(0x05,0x6F,0x00)
    r2 = p.add_run(text); _cn(r2); r2.font.color.rgb = RGBColor(0x33,0x33,0x33)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'E6F4EA')
    p._p.get_or_add_pPr().append(shd)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = p.paragraph_format.right_indent = Inches(0.4)
    r1 = p.add_run('注意：'); _cn(r1); r1.bold = True; r1.font.color.rgb = RGBColor(0x92,0x4E,0x00)
    r2 = p.add_run(text); _cn(r2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'FFF8E1')
    p._p.get_or_add_pPr().append(shd)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        r = cell.paragraphs[0].add_run(h); _cn(r); r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1E40AF')
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            r = cell.paragraphs[0].add_run(val); _cn(r)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def codeblock(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text); r.font.name = 'Courier New'; r.font.size = Pt(9)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F3F4F6')
    p._p.get_or_add_pPr().append(shd)

def pb(): doc.add_page_break()


# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('\n\n\n')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('照片管理系统'); _cn(r); r.font.size=Pt(32); r.bold=True; r.font.color.rgb=BLUE1
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Photo Management System  (PMS)'); _cn(r); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x60,0x7D,0x8B)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('用户手册'); _cn(r); r.font.size=Pt(28); r.bold=True; r.font.color.rgb=BLUE2
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('版本 1.0  |  ' + datetime.date.today().strftime('%Y年%m月'))
_cn(r); r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
pb()

# TOC
h1('目录')
toc = [
    ('1','快速入门'),('1.1','系统要求'),('1.2','下载与安装'),('1.3','首次启动'),
    ('2','界面介绍'),('2.1','整体布局'),('2.2','左侧标签栏'),('2.3','主内容区'),
    ('3','文件夹管理（树形视图）'),('3.1','新建文件夹'),('3.2','重命名与删除文件夹'),
    ('3.3','排序与移动文件夹'),('3.4','序列号（SN）'),
    ('4','文档操作'),('4.1','上传文件'),('4.2','上传文件夹'),('4.3','关联已有文档'),
    ('4.4','查看文档'),('4.5','下载与删除文档'),
    ('5','OCR 文字识别'),('5.1','OCR 工作原理'),('5.2','执行 OCR'),('5.3','保存 OCR 文本以供搜索'),
    ('6','关键参数提取'),('6.1','什么是关键参数'),('6.2','从文档中提取参数'),('6.3','管理全局关键参数'),
    ('7','搜索文档'),('7.1','基本搜索'),('7.2','使用筛选器'),('7.3','打开并定位搜索结果'),
    ('8','生成数据手册 PDF'),('8.1','选择文档'),('8.2','生成 PDF'),
    ('9','合并文档'),
    ('10','项目管理'),('10.1','导出项目（.dms）'),('10.2','导入项目（.dms）'),
    ('10.3','导出 CSV'),('10.4','批量 ZIP 导入'),('10.5','层级结构导入'),
    ('11','密码保护'),('12','键盘快捷键'),('13','常见问题与解决方法'),('','快速参考卡片'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4) if '.' in num else Inches(0)
    label = (num + '  ' + title) if num else title
    r = p.add_run(label); _cn(r)
    if '.' not in num and num: r.bold = True
pb()


# 1 快速入门
h1('1  快速入门')
h2('1.1  系统要求')
table(['项目','要求'],[
    ['操作系统','Windows 10 / 11（64 位）'],
    ['内存','最低 4 GB，推荐 8 GB'],
    ['磁盘空间','程序本身约 200 MB；另需空间存放您的文档'],
    ['显示器分辨率','1280 x 720 或更高'],
    ['网络连接','不需要（完全离线运行）。仅在查询 GPS 地址时需要联网。'],
],[1.8,4.2])

h2('1.2  下载与安装')
para('请按以下步骤在 Windows 电脑上安装并运行 PMS。')
numbered('从您的联系人处接收或下载 PMS.zip 文件（例如通过电子邮件、微信或 U 盘）。')
numbered('右键单击 PMS.zip，选择【全部解压缩】。')
numbered('选择目标文件夹，例如 C:\\PMS，然后单击【解压缩】。')
numbered('打开解压后的文件夹，找到 PMS.exe 文件。')
numbered('双击 PMS.exe 启动程序。')
tip('无需安装任何其他软件。PMS.exe 是独立程序，已内置 Python 及所有必要库。')
note('首次运行 PMS.exe 时，Windows 可能弹出 SmartScreen 安全警告。请单击【更多信息】然后单击【仍要运行】。此警告是因为该文件未经商业代码签名，程序本身是安全的。')
doc.add_paragraph()

h2('1.3  首次启动')
para('首次启动 PMS 时，请按以下步骤操作：')
numbered('后台会短暂出现一个控制台窗口，这是内置的 Web 服务器，请勿关闭它。')
numbered('默认浏览器会自动打开并访问 http://localhost:5000。')
numbered('您将看到 PMS 欢迎界面。')
numbered('系统会提示您选择存储文件夹——所有文件和数据都将保存在此处。')
bullet('单击顶部标题栏中的文件夹图标（或蓝色的【设置存储路径】按钮）。')
bullet('输入路径，例如 C:\\Users\\用户名\\Documents\\PMS_Data，然后按 Enter 键。')
bullet('如果该文件夹不存在，PMS 会自动创建。')
numbered('设置存储路径后，主界面加载完成，即可开始使用 PMS。')
tip('建议将存储路径设置在【文档】或【图片】文件夹内，以便纳入日常备份。')
pb()


# 2 界面介绍
h1('2  界面介绍')
h2('2.1  整体布局')
para('PMS 窗口分为三个主要区域：')
table(['区域','位置','用途'],[
    ['顶部标题栏','页面顶部','显示存储路径、文档总数、项目菜单及退出登录按钮。'],
    ['左侧边栏','左侧面板','包含树形视图、搜索、数据手册三个标签页。'],
    ['主内容区','右侧/中间区域','显示当前所选文件夹中的文档。'],
],[1.5,1.8,3.3])
tip('拖动左侧边栏与主内容区之间的分隔线可调整宽度；双击分隔线可恢复默认宽度。')

h2('2.2  左侧标签栏')
table(['标签','功能说明'],[
    ['树形','浏览和管理文件夹层级结构，是主要的导航视图。'],
    ['搜索','对所有文档进行全文和元数据搜索。'],
    ['数据手册','勾选文档后，生成带封面和目录的专业合并 PDF。'],
],[2.0,4.5])

h2('2.3  主内容区')
para('在树形标签中单击某个文件夹后，右侧面板将显示：')
bullet('文件夹名称、序列号（SN）和描述信息（位于顶部）。')
bullet('操作按钮：上传文件、上传文件夹、关联已有文档、合并文档。')
bullet('该文件夹中每个文档的卡片，显示文件名、大小、类型图标、上传日期及操作图标。')
pb()

# 3 文件夹管理
h1('3  文件夹管理（树形视图）')
para('PMS 使用树形结构组织文档，类似于 Windows 资源管理器，但额外支持序列号、描述和文档计数等功能。')

h2('3.1  新建文件夹')
numbered('在树形视图中单击某个文件夹以选中它（新文件夹将作为其子文件夹创建）。')
numbered('单击边栏工具栏中的新建文件夹按钮，或右键单击所选文件夹，选择【新建文件夹】。')
numbered('出现文本输入框后，输入文件夹名称，按 Enter 键确认，或按 Escape 键取消。')
tip('文件夹可以无限层级嵌套，没有深度限制。')

h2('3.2  重命名与删除文件夹')
para('重命名：')
bullet('双击文件夹名称，或右键单击然后选择重命名。')
bullet('编辑名称后按 Enter 键保存。')
para('删除：')
bullet('右键单击文件夹，选择删除。')
bullet('确认删除后，该文件夹及其所有子文件夹将从树形视图中移除。')
note('删除文件夹只会将其从树形视图中移除，磁盘上的实际文档文件仍保留在存储文件夹中。您仍可通过搜索标签或在 Windows 资源管理器中浏览存储文件夹找到这些文件。')

h2('3.3  排序与移动文件夹')
para('在同级文件夹间排序：')
bullet('将鼠标悬停在文件夹上，会出现向上和向下箭头按钮。')
bullet('单击向上箭头上移，单击向下箭头下移。')
para('移动到不同的上级文件夹：')
bullet('单击并拖动文件夹，将其拖放到目标父文件夹上（高亮显示放置目标）。')
bullet('也可以将其拖放到另一个文件夹的正上方或正下方，实现前置或后置排列。')

h2('3.4  序列号（SN）')
para('每个文件夹都可以设置可选的序列号（SN），用于标记组件、零件或项目编号。')
bullet('单击文件夹上的铅笔图标，打开其属性。')
bullet('在序列号字段中输入值（例如 SN-YT635R0267）。')
bullet('SN 将显示在树形视图中该文件夹名称旁边。')
bullet('文件夹内的文档会自动继承该文件夹的 SN，除非文档本身设置了独立的 SN。')
tip('如需为某个文档设置独立的 SN，请单击该文档卡片上的 SN 按钮。')
pb()


# 4 文档操作
h1('4  文档操作')
h2('4.1  上传文件')
numbered('在树形标签中选择目标文件夹。')
numbered('单击主内容区中的【上传文件】按钮。')
numbered('在文件选择对话框中选择一个或多个文件，然后单击【打开】。')
numbered('屏幕顶部出现进度条，显示正在上传的文件数量。')
numbered('上传完成后，新的文档卡片将出现在面板中。')
para('支持的文件类型：')
table(['类别','扩展名'],[
    ['图片','JPEG、PNG、GIF、WebP、BMP、TIFF、HEIC'],
    ['文档','PDF、DOC、DOCX、XLS、XLSX、PPT、PPTX、TXT、CSV、HTML、JSON、XML'],
    ['压缩包','ZIP'],
],[1.5,5.0])
tip('对于 JPEG 和 TIFF 格式的照片，PMS 会自动从 EXIF 元数据中读取拍摄日期和 GPS 位置，并与文档一起存储。')

h2('4.2  上传文件夹')
numbered('在树形视图中选择目标文件夹。')
numbered('单击【上传文件夹】按钮。')
numbered('在弹出的对话框中选择电脑上的一个文件夹。')
numbered('PMS 将递归上传所有文件，并在树形视图中创建对应的子文件夹。')

h2('4.3  关联已有文档')
para('同一个文档可以关联到多个文件夹，而不会在磁盘上复制文件。')
numbered('选择您想要关联文档的文件夹。')
numbered('单击【关联已有文档】。')
numbered('弹出对话框，列出系统中的所有文档。')
numbered('勾选要关联的文档，然后单击【添加关联】。')
numbered('该文档现在同时出现在原始文件夹和新文件夹中。')

h2('4.4  查看文档')
numbered('单击文档名称或缩略图，打开文档查看器。')
numbered('查看器以弹窗形式打开，左侧显示文档预览。')
bullet('图片直接内嵌显示。')
bullet('PDF 在内嵌查看器中打开，支持滚动和缩放。')
bullet('文本文件以等宽字体显示，方便阅读。')
bullet('Office 文件（Word、Excel、PowerPoint）提供下载按钮，可在本地 Office 应用中打开。')
numbered('查看器右侧有 OCR、关键参数提取和下载等操作按钮。')
numbered('单击 X 按钮或按 Escape 键关闭查看器。')

h2('4.5  下载与删除文档')
para('下载：单击文档卡片上的下载图标，或在查看器内单击【下载】按钮。')
para('从文件夹取消关联：单击文档卡片上的垃圾桶图标。如果该文档已关联到其他文件夹，则只从当前文件夹移除；文件和其他关联仍然保留。')
para('永久删除：只能从搜索标签或【未关联文档】列表中操作。永久删除会从所有文件夹中移除该文档，并从磁盘上删除文件。')
pb()

# 5 OCR
h1('5  OCR 文字识别')
h2('5.1  OCR 工作原理')
para('OCR（光学字符识别）将扫描图片和 PDF 转换为可搜索的文本。PMS 使用 Tesseract OCR 引擎，支持中文（简体）和英文。')
table(['方法','适用情况','速度'],[
    ['文本层提取','已包含数字文本层的 PDF','非常快'],
    ['Tesseract OCR','扫描版 PDF、图片或文本层损坏的文档','中等（取决于文件大小）'],
],[1.8,3.2,1.5])
note('如果服务器端未安装 Tesseract，则无法对图片进行 OCR。但对于含数字文本层的 PDF，文本层提取功能仍可正常使用。')

h2('5.2  执行 OCR')
numbered('单击文档名称，打开文档查看器。')
numbered('在右侧面板中单击【提取文字（OCR）】。')
numbered('PMS 开始提取文本，并在右侧面板中显示结果。')
numbered('如果 PDF 包含内嵌文本层，提取几乎是即时完成的。')
numbered('如果需要 OCR（扫描文档或图片），会显示进度指示。')
numbered('完成后，提取的文本显示在可编辑的文本框中，可直接编辑以修正 OCR 错误。')
numbered('面板还会显示：已处理页数、字符数、检测到的语言和使用的提取方法。')
note('如果 PDF 的文本层已损坏或内容不正确，请单击【强制 OCR】按钮，改用 Tesseract 重新提取。')

h2('5.3  保存 OCR 文本以供搜索')
numbered('查看并根据需要编辑提取的文本后，单击【保存并启用搜索】。')
numbered('文本将随文档元数据一起保存。')
numbered('此后，在搜索标签中搜索该文档包含的词语时，即可找到该文档。')
tip('每个文档只需执行一次 OCR。文本保存后，文档将永久可搜索。')
pb()


# 6 关键参数
h1('6  关键参数提取')
h2('6.1  什么是关键参数')
para('关键参数是用户自定义的标签（例如合同编号、压力等级、人物、事件），PMS 可以从文档文本中自动查找并提取对应的值。每个参数存储两个值：')
bullet('设计值——规定或要求的值（例如 3000 psi）。')
bullet('实际值——从文档中找到的实测或认证值。')
para('这样您就可以快速比较数百份文档的规格与实际值。')

h2('6.2  从文档中提取参数')
numbered('打开文档查看器。')
numbered('在右侧面板中单击【提取关键参数】。')
numbered('出现参数列表（全局默认参数加上文档专属参数）。')
numbered('勾选您想提取的参数。')
numbered('可选：切换【仅同行匹配】开关，仅在参数名称所在行内匹配值。')
numbered('单击【提取】。')
numbered('PMS 在文档文本中搜索，并在结果表中填入数值：')
table(['列','含义'],[
    ['参数','正在搜索的标签名'],
    ['找到的值','PMS 在文档文本中找到的值'],
    ['状态','已找到，或 NF（未找到）'],
],[1.5,5.0])
numbered('如需修改，可直接编辑任意值。')
numbered('单击【保存参数】，将这些值存储到文档的元数据中。')
tip('元数据是全局共享的——修改后，所有引用此文档的文件夹均会同步更新。')

h2('6.3  管理全局关键参数')
numbered('单击顶部标题栏中的【项目】菜单。')
numbered('选择【全局关键参数】。')
numbered('在弹出的对话框中，可以添加、删除或调整参数顺序。')
numbered('单击【恢复默认值】可重置为默认的参数列表。')
numbered('更改立即生效，适用于后续所有提取操作。')
pb()

# 7 搜索
h1('7  搜索文档')
h2('7.1  基本搜索')
numbered('单击左侧边栏中的【搜索】标签。')
numbered('在搜索框中输入关键词或短语。')
numbered('搜索结果即时以卡片形式显示，每张卡片包含：')
bullet('文档名称和文件类型图标。')
bullet('上传日期、文件大小和文档 ID。')
bullet('所有元数据字段（参数名称 / 设计值 / 实际值）。')
bullet('GPS 定位地址（如有）。')
bullet('引用此文档的文件夹列表。')
note('搜索不区分大小写，支持部分匹配——例如输入【泵】可以找到【离心泵】【水泵总成】等。')

h2('7.2  使用筛选器')
table(['筛选器','使用方法'],[
    ['文件类型','勾选图片、PDF、文本或文档，将结果限定为该类型。'],
    ['范围节点','启用【限定范围至所选节点】，仅在当前文件夹及其子文件夹中搜索。'],
    ['元数据','输入参数名称和/或值，查找具有匹配元数据的文档（例如：压力 = 3000 psi）。'],
],[1.5,5.0])

h2('7.3  打开并定位搜索结果')
bullet('单击结果卡片，打开该文档的查看器。')
bullet('双击结果卡片，切换到树形标签，并直接跳转到包含该文档的文件夹。')
pb()

# 8 数据手册
h1('8  生成数据手册 PDF')
para('数据手册功能可将选定的文档合并为一份专业 PDF，包含封面、目录、书签、章节标题和页码。')
h2('8.1  选择文档')
numbered('单击左侧边栏中的【数据手册】标签。')
numbered('树形视图中每个文档旁边会出现复选框。')
numbered('勾选单个文档，或使用以下快捷操作：')
bullet('单击【选择该节点所有文档】，选中特定文件夹中的所有文档。')
bullet('单击【选择子树所有文档】，选中文件夹及其所有子文件夹中的文档。')
bullet('单击【清除】取消所有选择。')
tip('您的选择会自动保存——关闭并重新打开 PMS 后，选择状态仍然保留。')

h2('8.2  生成 PDF')
numbered('选好文档后，单击【生成数据手册】。')
numbered('在弹出的对话框中填写标题和副标题。')
bullet('标题——封面上显示的主标题。')
bullet('副标题——辅助说明行（例如项目名称、版本号）。')
numbered('单击【生成】。')
numbered('PMS 在服务器端生成 PDF，您的浏览器将自动下载该文件。')
para('生成的 PDF 包含以下内容：封面（标题、副标题和日期）、带可点击书签的目录、章节标题页、按树形顺序排列的所有选定文档、每页均有页码和页脚。')
pb()


# 9 合并文档
h1('9  合并文档')
para('合并功能可将一个文件夹中的多个 PDF（及图片）合并为单个 PDF 文件，并将结果保存回该文件夹。')
numbered('在树形标签中选择目标文件夹。')
numbered('单击主内容区中的【合并文档】按钮。')
numbered('弹出对话框，列出该文件夹中的所有文档。')
numbered('可选：勾选【包含子文件夹】，同时合并子文件夹中的文档。')
numbered('勾选要合并的文档（或单击【全选】）。')
numbered('在【输出文件名】字段中输入合并后文件的名称。')
numbered('单击【合并】。')
numbered('PMS 创建合并后的 PDF，并将其添加到当前文件夹中。')
note('合并后，原始文档仍保留在文件夹中。如不再需要单独的副本，可手动取消关联。')
pb()

# 10 项目管理
h1('10  项目管理')
para('顶部标题栏中的【项目】菜单提供备份、恢复和导入数据的工具。')

h2('10.1  导出项目（.dms）')
para('将整个项目完整备份——包括所有文件夹结构、文档元数据和文档文件——保存为单个 .dms 文件（本质上是 ZIP 压缩包）。')
numbered('单击标题栏中的【项目】，选择【导出项目（.dms）】。')
numbered('浏览器将下载备份文件（例如 project_backup_20260610.dms）。')
tip('请将此文件保存到 U 盘或云存储，作为定期备份。')

h2('10.2  导入项目（.dms）')
para('在任何运行 PMS 的电脑上还原之前导出的项目。')
numbered('单击【项目】→【导入项目（.dms）】。')
numbered('选择 .dms 文件。')
numbered('在对话框中选择电脑上的目标文件夹（建议选择新的空文件夹）。')
numbered('单击【导入】。PMS 解压所有文件并重建树形结构。')
note('如果导入到已有数据的文件夹，系统会提示您确认是否覆盖。')

h2('10.3  导出 CSV')
para('将所有文档的元数据导出为 CSV 电子表格。')
numbered('单击【项目】→【导出 CSV】。浏览器下载 CSV 文件。')
para('CSV 文件每行对应一个文档，列信息包括：文档 ID、名称、文件大小、上传日期及所有元数据参数字段。')
tip('在 Microsoft Excel 或 Google 表格中打开 CSV，可进行数据分析、报告生成或与同事共享。')

h2('10.4  批量 ZIP 导入')
para('通过命名规范快速上传大量文件，并自动分配到对应文件夹。')
numbered('准备一个 ZIP 文件，按以下格式命名每个文件：文件夹名#说明.pdf（例如：泵总成#测试报告.pdf）。')
numbered('单击【项目】→【批量 ZIP 导入】，选择您的 ZIP 文件。')
numbered('PMS 读取每个文件名，将 # 号前的部分与树形视图中的文件夹名称匹配，并将文档放入对应文件夹。')
note('导入前，文件夹必须已在树形视图中存在。文件夹名称无法匹配的文档将进入未关联状态。')

h2('10.5  层级结构导入（从文本文件创建文件夹树）')
para('通过纯文本文件快速创建文件夹树结构。')
numbered('准备一个文本文件（.txt），每行格式为：节点ID [父节点ID]')
para('示例：')
codeblock('总成\n泵壳组件 总成\n转子组件 总成\n叶轮 转子组件\n主轴 转子组件')
numbered('单击【项目】→【从层级结构创建文件夹树】。')
numbered('上传您的文本文件（或粘贴文本内容）。')
numbered('PMS 验证结构，显示预览，并报告将创建的节点数量。')
numbered('单击【创建】，所有文件夹立即出现在树形视图中。')
pb()

# 11 密码
h1('11  密码保护')
para('您可以为 PMS 设置密码，确保只有授权用户才能访问数据。')
para('设置密码：')
numbered('单击【项目】→【设置密码】。')
numbered('输入并确认您选择的密码，然后单击【设置】。')
numbered('下次启动 PMS 时，在主界面加载前会显示登录界面。')
para('取消密码：')
numbered('单击【项目】→【修改密码】。')
numbered('输入当前密码以确认身份，然后将新密码字段留空，单击【设置】即可取消密码保护。')
para('退出登录：单击标题栏中的【退出登录】按钮。PMS 返回登录界面；服务器继续运行，仅浏览器会话结束。')
note('密码以加盐哈希方式存储在服务器端（SHA-256，100,000 次迭代），不以明文保存。')
pb()

# 12 快捷键
h1('12  键盘快捷键')
table(['操作','快捷键'],[
    ['确认文件夹重命名 / 新文件夹名称','Enter'],
    ['取消重命名 / 新建文件夹','Escape'],
    ['开始重命名文件夹','双击文件夹名称'],
    ['打开文件夹右键菜单','右键单击文件夹'],
    ['关闭文档查看器','Escape 或单击 X'],
    ['复制选中文本','Ctrl + C'],
    ['粘贴文本','Ctrl + V'],
    ['全选文本框内容','Ctrl + A'],
    ['撤销（文本框内）','Ctrl + Z'],
],[3.5,3.0])
pb()

# 13 故障排除
h1('13  常见问题与解决方法')
table(['问题','解决方法'],[
    ['Windows SmartScreen 阻止 PMS.exe 运行',
     '单击【更多信息】，然后单击【仍要运行】。此提示是因为文件未经商业签名，程序本身是安全的。'],
    ['浏览器未自动打开',
     '手动打开浏览器，访问 http://localhost:5000。'],
    ['浏览器显示无法访问此网站',
     '请确认后台的 PMS.exe 控制台窗口仍在运行。如已关闭，请重新启动 PMS.exe。'],
    ['OCR 按钮显示为灰色或提示 Tesseract 不可用',
     'Tesseract OCR 是可选组件，需单独安装。请从 Tesseract 官方网站下载安装，并重启 PMS。'],
    ['中文文字未能正确识别',
     '请安装 Tesseract 的简体中文语言包（chi_sim.traineddata），详见 Tesseract 文档。'],
    ['GPS 位置显示为坐标而非城市名称',
     '地址反查服务需要网络连接，请检查您的网络状态。'],
    ['上传的文件未出现在文档列表中',
     '请刷新浏览器页面（F5）。文件可能已成功上传，但页面未自动更新。'],
    ['找不到之前上传的文档',
     '请使用搜索标签，按文件名或内容搜索。该文档可能被放在了不同的文件夹中。'],
    ['PMS.exe 启动后立即关闭',
     '请确保电脑上的 5000 端口未被其他程序占用，重启电脑后再试。'],
    ['生成数据手册 PDF 失败',
     '请确保所有选定文档为有效的 PDF 或图片文件。损坏的文件可能导致生成失败。'],
],[2.5,4.0])
doc.add_paragraph()
h2('获取帮助')
para('如遇到上述列表中未包含的问题，请联系应用开发者并提供以下信息：')
bullet('问题发生时您正在执行的操作描述。')
bullet('屏幕上显示的任何错误信息。')
bullet('您的操作系统版本（例如 Windows 11）。')
pb()

# 快速参考卡片
h1('快速参考卡片')
table(['任务','操作方法'],[
    ['新建文件夹','选择父文件夹 → 单击新建文件夹 → 输入名称 → Enter'],
    ['上传文档','选择文件夹 → 上传文件 → 选择文件'],
    ['查看文档','单击文档名称'],
    ['对文档执行 OCR','打开查看器 → 提取文字（OCR）→ 保存并启用搜索'],
    ['搜索所有文档','单击搜索标签 → 输入关键词'],
    ['导出项目备份','项目菜单 → 导出项目（.dms）'],
    ['生成合并 PDF','数据手册标签 → 勾选文档 → 生成数据手册'],
    ['合并文件为单个 PDF','选择文件夹 → 合并文档 → 选择文件 → 合并'],
    ['移动文件夹','拖放文件夹到树形视图中的新位置'],
    ['设置序列号','单击文件夹铅笔图标 → 输入 SN → 完成'],
    ['导出元数据为 CSV','项目菜单 → 导出 CSV'],
    ['设置访问密码','项目菜单 → 设置密码'],
],[2.5,4.5])

out = '/Users/david/PMS/PMS用户手册.docx'
doc.save(out)
print('已保存：' + out)
